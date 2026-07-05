"""DOCX export rendering (C8)."""

from __future__ import annotations

from core.export.stages import meeting_stage_heading, report_stage_heading, stage_heading
from core.export.types import ExportPayload


def docx_dependency_available() -> bool:
    try:
        import docx  # noqa: F401

        return True
    except ImportError:
        return False


def render_docx_document(
    payload: ExportPayload,
    *,
    stages: tuple[int, ...],
    include_meta: bool,
    template_id: str,
):
    from docx import Document

    document = Document()
    if template_id == "meeting":
        _render_meeting(document, payload, stages=stages, include_meta=include_meta)
    elif template_id == "report":
        _render_report(document, payload, stages=stages, include_meta=include_meta)
    else:
        _render_basic(document, payload, stages=stages, include_meta=include_meta)
    return document


def _render_basic(document, payload: ExportPayload, *, stages: tuple[int, ...], include_meta: bool) -> None:
    document.add_heading(payload.mode_name, level=0)
    if include_meta:
        document.add_paragraph(f"생성: {payload.created_at.strftime('%Y-%m-%d %H:%M')}")
        document.add_paragraph(f"세션: {payload.session_id}")
    for stage in stages:
        if stage not in payload.stages:
            continue
        document.add_heading(stage_heading(stage), level=2)
        _add_paragraphs(document, payload.stages[stage])
        _append_provider(document, payload.providers.get(stage))


def _render_meeting(document, payload: ExportPayload, *, stages: tuple[int, ...], include_meta: bool) -> None:
    document.add_heading("회의록", level=0)
    document.add_paragraph(f"일시: {payload.created_at.strftime('%Y-%m-%d %H:%M')}")
    document.add_paragraph(f"모드: {payload.mode_name}")
    if include_meta:
        document.add_paragraph(f"세션 ID: {payload.session_id}")
    for stage in stages:
        if stage not in payload.stages:
            continue
        document.add_heading(meeting_stage_heading(stage), level=1)
        _add_paragraphs(document, payload.stages[stage])
        _append_provider(document, payload.providers.get(stage))


def _render_report(document, payload: ExportPayload, *, stages: tuple[int, ...], include_meta: bool) -> None:
    document.add_heading("보고서", level=0)
    document.add_paragraph(f"작성일: {payload.created_at.strftime('%Y-%m-%d %H:%M')}")
    document.add_paragraph(f"주제: {payload.mode_name}")
    if include_meta:
        document.add_paragraph(f"세션 ID: {payload.session_id}")
    for stage in stages:
        if stage not in payload.stages:
            continue
        document.add_heading(report_stage_heading(stage), level=1)
        _add_paragraphs(document, payload.stages[stage])
        _append_provider(document, payload.providers.get(stage))


def _append_provider(document, provider: str | None) -> None:
    if provider:
        document.add_paragraph(f"provider: {provider}")


def _add_paragraphs(document, text: str) -> None:
    blocks = [block.strip() for block in text.split("\n\n") if block.strip()]
    if not blocks:
        document.add_paragraph("")
        return
    for block in blocks:
        lines = block.splitlines()
        if all(line.strip().startswith(("- ", "* ")) for line in lines if line.strip()):
            for line in lines:
                stripped = line.strip()
                if stripped.startswith(("- ", "* ")):
                    document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(block)
